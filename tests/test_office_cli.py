import io

import fitz
import pytest
from docx import Document

from office.convert import convert_to_markdown
from office.cli import main


def _docx(path):
    doc = Document()
    doc.add_heading("Hello Doc", level=1)
    doc.add_paragraph("body text here")
    doc.save(path)


def test_dispatch_docx(tmp_path):
    p = str(tmp_path / "d.docx")
    _docx(p)
    result = convert_to_markdown(p)
    assert "# Hello Doc" in result.markdown


def test_dispatch_bad_extension(tmp_path):
    p = str(tmp_path / "x.txt")
    open(p, "w").close()
    with pytest.raises(ValueError):
        convert_to_markdown(p)


def test_dispatch_missing_file():
    with pytest.raises(FileNotFoundError):
        convert_to_markdown("nope.docx")


def test_cli_writes_markdown(tmp_path, capsys):
    p = str(tmp_path / "d.docx")
    _docx(p)
    out = str(tmp_path / "out.md")
    rc = main([p, "-o", out])
    assert rc == 0
    with open(out, encoding="utf-8") as f:
        content = f.read()
    assert "# Hello Doc" in content
    err = capsys.readouterr().err
    assert "Format:" in err and "units:" in err and "images:" in err


def test_cli_bad_extension_returns_error(tmp_path, capsys):
    p = str(tmp_path / "x.txt")
    open(p, "w").close()
    rc = main([p])
    assert rc == 1
    assert "Error" in capsys.readouterr().err


def test_cli_stdout(tmp_path, capsys):
    p = str(tmp_path / "d.docx")
    _docx(p)
    rc = main([p])
    assert rc == 0
    out = capsys.readouterr()
    assert "# Hello Doc" in out.out
    assert "Format:" in out.err


def test_cli_math_rules(capsys):
    from office.cli import main
    rc = main(["--math-rules"])
    assert rc == 0
    out = capsys.readouterr().out
    assert "[(" in out and "sqrt" in out
