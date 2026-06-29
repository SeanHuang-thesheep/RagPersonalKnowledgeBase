import io

import fitz  # 已装，用来生成合法 PNG
from docx import Document

from office.docx_convert import docx_to_markdown


def _png():
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 3, 3))
    pix.set_rect(pix.irect, (10, 20, 30))
    return pix.tobytes("png")


def _build(path):
    doc = Document()
    doc.add_heading("Title One", level=1)
    doc.add_heading("Sub Two", level=2)
    doc.add_paragraph("plain body paragraph")
    doc.add_paragraph("bullet item", style="List Bullet")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "h1"
    table.cell(0, 1).text = "h2"
    table.cell(1, 0).text = "a"
    table.cell(1, 1).text = "b"
    doc.add_picture(io.BytesIO(_png()))
    doc.save(path)


def test_docx_basic(tmp_path):
    p = str(tmp_path / "d.docx")
    _build(p)
    result = docx_to_markdown(p, asset_dir=str(tmp_path / "assets"))
    md = result.markdown
    assert "# Title One" in md
    assert "## Sub Two" in md
    assert "plain body paragraph" in md
    assert "- bullet item" in md
    assert "| h1 | h2 |" in md
    assert "| --- | --- |" in md
    assert "![](" in md
    assert result.image_count == 1
    assert result.unit_count >= 5


def test_docx_missing_file():
    import pytest
    with pytest.raises(FileNotFoundError):
        docx_to_markdown("nope.docx")


def test_docx_equation_placeholder(tmp_path):
    from docx.oxml import parse_xml
    doc = Document()
    para = doc.add_paragraph("")
    omath = parse_xml(
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '<m:r><m:t>x</m:t></m:r></m:oMath>'
    )
    para._p.append(omath)
    p = str(tmp_path / "eq.docx")
    doc.save(p)
    md = docx_to_markdown(p).markdown
    assert "$x$" in md


def test_docx_two_distinct_images(tmp_path):
    def _png(color):
        pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 3, 3))
        pix.set_rect(pix.irect, color)
        return pix.tobytes("png")

    doc = Document()
    doc.add_paragraph("first")
    doc.add_picture(io.BytesIO(_png((10, 20, 30))))
    doc.add_paragraph("second")
    doc.add_picture(io.BytesIO(_png((200, 100, 50))))
    p = str(tmp_path / "imgs.docx")
    doc.save(p)
    result = docx_to_markdown(p, asset_dir=str(tmp_path / "assets"))
    assert result.image_count == 2
    assert len(set(result.images)) == 2   # 两个不同文件，无撞名覆盖


def test_docx_omml_linearized(tmp_path):
    from docx import Document
    from docx.oxml import parse_xml

    M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    doc = Document()
    para = doc.add_paragraph("")
    omath = parse_xml(
        f'<m:oMath xmlns:m="{M}">'
        f'<m:f><m:num><m:r><m:t>dy</m:t></m:r></m:num>'
        f'<m:den><m:r><m:t>dt</m:t></m:r></m:den></m:f>'
        f'<m:r><m:t> + </m:t></m:r>'
        f'<m:sSup><m:e><m:r><m:t>y</m:t></m:r></m:e>'
        f'<m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>'
        f'<m:r><m:t> = 0</m:t></m:r>'
        f'</m:oMath>'
    )
    para._p.append(omath)
    p = str(tmp_path / "eq.docx")
    doc.save(p)

    from office.docx_convert import docx_to_markdown
    md = docx_to_markdown(p).markdown
    assert "$(dy)/(dt) + y^2 = 0$" in md
    assert "[equation]" not in md
